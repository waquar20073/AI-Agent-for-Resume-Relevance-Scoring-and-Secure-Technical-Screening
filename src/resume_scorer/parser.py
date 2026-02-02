import PyPDF2
import docx
import pdfplumber
from typing import Dict, List, Any
from datetime import datetime
import uuid
import logging
from src.utils.nlp_utils import NLPProcessor
from src.utils.data_models import ResumeData, Skill, JobDescription

logger = logging.getLogger(__name__)

class ResumeParser:
    def __init__(self):
        self.nlp_processor = NLPProcessor()
    
    def parse_resume(self, file_path: str, file_content: bytes = None, filename: str = None) -> ResumeData:
        """Parse resume from file and extract structured data."""
        try:
            if file_content:
                text = self._extract_text_from_bytes(file_content, filename)
            else:
                text = self._extract_text_from_file(file_path)
            
            # Clean and process text
            cleaned_text = self.nlp_processor.extract_text_from_resume(text)
            
            # Extract components
            skills_data = self.nlp_processor.extract_skills(cleaned_text)
            experience_years = self.nlp_processor.extract_experience_years(cleaned_text)
            education_data = self.nlp_processor.extract_education(cleaned_text)
            
            # Convert to skill objects
            skills = self._convert_to_skill_objects(skills_data, cleaned_text)
            
            # Extract certifications
            certifications = self._extract_certifications(cleaned_text)
            
            resume_data = ResumeData(
                candidate_id=str(uuid.uuid4()),
                skills=skills,
                experience_years=experience_years,
                education=education_data,
                certifications=certifications,
                raw_text=cleaned_text,
                processed_at=datetime.now()
            )
            
            logger.info(f"Successfully parsed resume for candidate {resume_data.candidate_id}")
            return resume_data
            
        except Exception as e:
            logger.error(f"Error parsing resume: {e}")
            raise
    
    def parse_job_description(self, job_text: str, job_title: str = None) -> JobDescription:
        """Parse job description and extract requirements."""
        try:
            cleaned_text = self.nlp_processor.extract_text_from_resume(job_text)
            
            # Extract requirements
            skills_data = self.nlp_processor.extract_skills(cleaned_text)
            experience_required = self.nlp_processor.extract_experience_years(cleaned_text)
            
            # Convert to skill objects
            required_skills = self._convert_to_skill_objects(skills_data, cleaned_text)
            
            # Extract education requirements
            education_requirements = self._extract_education_requirements(cleaned_text)
            
            # Extract responsibilities
            responsibilities = self._extract_responsibilities(cleaned_text)
            
            job_description = JobDescription(
                job_id=str(uuid.uuid4()),
                title=job_title or "Unknown Position",
                required_skills=required_skills,
                experience_required=experience_required,
                education_requirements=education_requirements,
                responsibilities=responsibilities,
                raw_text=cleaned_text
            )
            
            logger.info(f"Successfully parsed job description for {job_description.job_id}")
            return job_description
            
        except Exception as e:
            logger.error(f"Error parsing job description: {e}")
            raise
    
    def _extract_text_from_file(self, file_path: str) -> str:
        """Extract text from various file formats."""
        file_extension = file_path.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            return self._extract_from_pdf(file_path)
        elif file_extension in ['docx', 'doc']:
            return self._extract_from_docx(file_path)
        elif file_extension == 'txt':
            return self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _extract_text_from_bytes(self, file_content: bytes, filename: str) -> str:
        """Extract text from file content bytes."""
        if not filename:
            raise ValueError("Filename is required when processing file content")
        
        file_extension = filename.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            return self._extract_from_pdf_bytes(file_content)
        elif file_extension in ['docx', 'doc']:
            return self._extract_from_docx_bytes(file_content)
        elif file_extension == 'txt':
            return file_content.decode('utf-8', errors='ignore')
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() or ""
            except Exception as e2:
                logger.error(f"Both PDF extraction methods failed: {e2}")
                raise
        return text
    
    def _extract_from_pdf_bytes(self, file_content: bytes) -> str:
        """Extract text from PDF bytes."""
        text = ""
        try:
            import io
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
            try:
                import io
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                for page in pdf_reader.pages:
                    text += page.extract_text() or ""
            except Exception as e2:
                logger.error(f"Both PDF extraction methods failed: {e2}")
                raise
        return text
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _extract_from_docx_bytes(self, file_content: bytes) -> str:
        """Extract text from DOCX bytes."""
        import io
        doc = docx.Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()
    
    def _convert_to_skill_objects(self, skills_data: Dict[str, List[str]], text: str) -> List[Skill]:
        """Convert skills dictionary to Skill objects."""
        skills = []
        for category, skill_list in skills_data.items():
            for skill_name in skill_list:
                proficiency = self.nlp_processor.assess_skill_proficiency(text, skill_name)
                # Estimate years of experience based on context
                years_exp = self._estimate_skill_experience(skill_name, text)
                
                skill = Skill(
                    name=skill_name,
                    category=category,
                    proficiency_level=proficiency,
                    years_experience=years_exp
                )
                skills.append(skill)
        
        return skills
    
    def _estimate_skill_experience(self, skill: str, text: str) -> float:
        """Estimate years of experience for a specific skill."""
        text_lower = text.lower()
        skill_lower = skill.lower()
        
        # Look for patterns indicating experience with specific skill
        patterns = [
            rf'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of\s*)?{re.escape(skill_lower)}',
            rf'{re.escape(skill_lower)}\s*(?:for\s*)?(\d+)\+?\s*(?:years?|yrs?)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text_lower)
            if match:
                return float(match.group(1))
        
        return 0.0
    
    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications from text."""
        certification_keywords = [
            'certified', 'certification', 'certificate', 'aws certified', 'google cloud certified',
            'microsoft certified', 'oracle certified', 'pmp', 'cfa', 'cpa', 'phd', 'master',
            'bachelor', 'mba', 'msc', 'bsc'
        ]
        
        certifications = []
        text_lower = text.lower()
        
        for keyword in certification_keywords:
            if keyword in text_lower:
                # Extract context around the keyword
                index = text_lower.find(keyword)
                start = max(0, index - 20)
                end = min(len(text), index + 50)
                context = text[start:end].strip()
                certifications.append(context)
        
        return list(set(certifications))
    
    def _extract_education_requirements(self, text: str) -> List[str]:
        """Extract education requirements from job description."""
        education_patterns = [
            r'(bachelor|master|phd|doctorate|mba|b\.s\.|m\.s\.|b\.a\.|m\.a\.)',
            r'(degree|graduation|diploma)',
            r'(computer science|data science|engineering|mathematics|statistics)'
        ]
        
        requirements = []
        text_lower = text.lower()
        
        for pattern in education_patterns:
            matches = re.findall(pattern, text_lower)
            requirements.extend(matches)
        
        return list(set(requirements))
    
    def _extract_responsibilities(self, text: str) -> List[str]:
        """Extract job responsibilities from description."""
        # Look for bullet points or numbered lists
        responsibilities = []
        
        # Split by common bullet point indicators
        bullet_patterns = [r'•', r'·', r'-', r'\*', r'\d+\.', r'\)']
        
        for pattern in bullet_patterns:
            parts = re.split(pattern, text)
            for part in parts[1:]:  # Skip the first part (before first bullet)
                cleaned = part.strip()
                if len(cleaned) > 10 and len(cleaned) < 200:
                    responsibilities.append(cleaned)
        
        # If no bullets found, split by sentences and look for responsibility keywords
        if not responsibilities:
            sentences = re.split(r'[.!?]+', text)
            responsibility_keywords = ['develop', 'design', 'implement', 'manage', 'lead', 'create', 'analyze']
            
            for sentence in sentences:
                sentence = sentence.strip()
                if any(keyword in sentence.lower() for keyword in responsibility_keywords):
                    if len(sentence) > 10:
                        responsibilities.append(sentence)
        
        return responsibilities[:10]  # Limit to top 10 responsibilities
